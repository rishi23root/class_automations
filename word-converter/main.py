# python main.py
from docx import Document # pip install python-docx   
from docx.shared import Inches,Pt,RGBColor
from PIL import Image
import json
import os


class project_creator:
    def __init__(self,info_file):
        # file name saved 
        self.info_file = info_file

    def __enter__(self):
        self.extract_info()
        # start a new word file 
        self.document = Document()
        # use sections to edit page layout and margin
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        return self
    
    def __exit__(self ,type, value, traceback):
        # save the file at last
        self.document.save(self.file_name +'.docx')
        print(self.file_name," Saved 😁😀")
        return False

    def extract_info(self):
        # data file and read data and update class
        with open(self.info_file,'r',encoding="utf8") as info_file:
            info = json.load(info_file)
            self.file_name = info["file_name"]
            self.questions = info["data"]["questions"]
            self.files = info["data"]["files"]
    
    def file_reader(self,file_name):
        print(f'reading file => {file_name} 🕵️‍')
        # try except block for error 
        try:
            with open(file_name ,'r') as code:
                return code.read()
        except :
            raise Exception("file reading error -check file name and location ")


    def question(self,data):
        # add heading for questions
        que = self.document.add_heading(data, level=1)
        font = que.style.font
        font.size = Pt(20)
        font.color.rgb = RGBColor(0,0,0)
    
    def code(self,data):
        # sol
        ans = self.document.add_paragraph("")
        ans.add_run('Sol.').bold = True

        # add code to the doc
        code = self.document.add_paragraph(data)
        # indent for code
        paragraph_format = code.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        font = code.style.font
        font.size = Pt(14)
        font.color.rgb = RGBColor(0,0,50)
    
    def image(self,image):
        # add iamge to the doc
        # output
        output = self.document.add_paragraph("")
        output.add_run('output.').bold = True
        
        # image
        # width and height of image
        w,_ = Image.open(image).size
        # big - 1920 1080 
        # idle - 700 _

        # print(w)
        if w < 700:
            self.document.add_picture(image)
        else:
            self.document.add_picture(image,width=Inches(7.5))


    def create_file(self):
        # write file name at top in center
        heading = self.document.add_heading(self.file_name, 0)  # print(self.file_name)
        heading.alignment = 1

        for question,files in zip(self.questions,self.files):
            # write questions as heading in bold
            # print(question)
            self.question(question)

            # open and read code file and butify it 
            code = self.file_reader(files[0])
            self.code(code)  # print(code)
            
            # add image(screen shot) in the end 
            ss = files[1]
            self.image(ss)   # print(ss)
            
            # add page break for every new question
            if question != self.questions[-1]:
                self.document.add_page_break()

    @staticmethod
    def open_file(file_name):
        # to open the fiel for results 
        os.startfile(file_name+'.docx')

    @classmethod
    def runner(cls,info_file):
        with cls(info_file) as clas :
            # this will create file and save all the changes
            clas.create_file()
            file_name = clas.file_name
        
        clas.open_file(file_name)


if __name__ == "__main__":
    project_creator.runner("info.json")