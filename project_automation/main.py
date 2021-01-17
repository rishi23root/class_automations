from docx import Document # pip install python-docx   
from docx.shared import Inches,Pt,RGBColor
from PIL import Image
import json
import os
import math
import argparse

class project_creator:
    max_img_width = 700
    def __init__(self,info_file,new_page_new_question = False):
        # file name saved 
        self.info_file = info_file
        self.new_page_new_question = new_page_new_question 

    def __enter__(self):
        self.extract_info()
        # start a new word file 
        self.document = Document()
        # use sections to edit page layout and margin
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        return self
    
    def __exit__(self ,type, value, traceback):
        # save the file at last
        self.document.save(self.file_name +'.docx')
        print(self.file_name," Saved 😁😀")
        return False

    def extract_info(self):
        # data file and read data and update class
        with open(self.info_file,'r',encoding="utf8") as info_file:
            try:
                info = json.load(info_file)
                self.file_name = info["file_name"]
                self.userinfo = info["userinfo"]
                self.questions = info["data"]["questions"]
                self.files = info["data"]["files"]
                self.isHeading = info["format"]["heading"]
                self.block = info["format"]["data"]
                self.isEnd_name = info["format"]["end_name"]            
            except :
                raise Exception(f' Unable to read data from json File -> {self.info_file}  -try fixing some format.')
            else:
                # for the directory for code should be present
                directory = os.path.normpath(info["directory"])
                if os.path.isdir(directory):
                    self.directory = directory 
                else :
                    raise Exception('Directory not found') 
           
    def file_reader(self,file_name):
        print(f'reading file => {file_name} 🕵️‍')
        # try except block for error 
        try:
            with open(file_name ,mode = 'r',encoding="utf8") as code:
                return code.read()
        except :
            raise Exception("file reading error -check file name and location (we are using encoding='utf8' for decoding)")

    def question(self,data):
        # add heading for questions
        que = self.document.add_heading(data, level=1)
        font = que.style.font
        font.size = Pt(20)
        font.color.rgb = RGBColor(0,0,0)
    
    def code(self,data):
        # sol
        ans = self.document.add_paragraph("")
        ans.add_run('Sol.').bold = True

        # add code to the doc
        code = self.document.add_paragraph(data)
        # indent for code
        paragraph_format = code.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        font = code.style.font
        font.size = Pt(14)
        font.color.rgb = RGBColor(0,0,50)
    
    def image(self,image):
        # add iamge to the doc
        # output
        output = self.document.add_paragraph("")
        output.add_run('output.').bold = True
        
        # image width and height of image
        w,h = Image.open(image).size

        # big - 1920 1080  # idle - 700 _
        if w < self.max_img_width:
            self.document.add_picture(image)
        else:
            self.document.add_picture(image,width=Inches(7.5))
        
    def data_block(self,question,code_data,ss):
        # write questions as heading in bold
        if self.block["question"] :
            self.question(question) # print(question)

        if self.block["solution"] :
            self.code(code_data)  # print(code_data)
           
        if self.block["picture"] : 
            self.image(ss)   # print(ss)
             
    def create_file(self):
        # write file name at top in center
        if self.isHeading:
            heading = self.document.add_heading(self.file_name, 0)  # print(self.file_name)
            heading.alignment = 1

        for question,files in zip(self.questions,self.files):
            # question  # code = self.file_reader(files[0]) # ss = files[1]
            self.data_block(question,self.file_reader(files[0]),files[1])

            if self.new_page_new_question and question != self.questions[-1]:
                self.document.add_page_break()

        if self.isEnd_name:
            # add name of the student in the end the file 
            hr = self.document.add_paragraph("")
            hr.add_run("_____________________________________________________________________").bold = True
            hr.alignment = 1
            para = ''
            for key,value in zip(self.userinfo.keys(),self.userinfo.values()):
                para += f'{key} : {value}\n' 
            else : para = para[:-1]

            end_userinfo = self.document.add_paragraph(para)
            font = end_userinfo.style.font
            font.color.rgb = RGBColor(0,0,0)

    def walk_in_dir(self):
        print(self.directory)
        os.chdir(self.directory)
        print(f"Dir changed to {self.directory}")

    @staticmethod
    def open_file(file_name):
        # to open the file for results 
        print("opening file for checking ...")
        os.startfile(file_name+'.docx')

    @classmethod
    def runner(cls,info_file,new_page_new_question = False):
        with cls(info_file,new_page_new_question) as clas:
            # change dir to the directory for reading and saving the data
            clas.walk_in_dir()
            
            # this will create file and save all the changes
            clas.create_file()
            file_name = clas.file_name
        
        # to open the after it saved
        cls.open_file(file_name)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('-n','--newPage',action="store_true",help = "for get every new question on new page.")
    parser.add_argument('-i','--infofile', default = 'info.json', help = "json data file info-file default (info.json).")
    
    args = parser.parse_args()
    
    project_creator.runner(args.infofile,new_page_new_question = args.newPage)
